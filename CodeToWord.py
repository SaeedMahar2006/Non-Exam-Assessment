#convenience tool for passting code into word doc.
from docx import Document
import os



def add_heading(text):
    global document
    document.add_heading(text, level=2)

def read_code(direct):
    global document
    f=open(direct,"r", encoding="utf8")
    text=f.read()
    document.add_paragraph(text)

def save_doc():
    global document

    document.save('code.docx')

if __name__=="__main__":
    document = Document()

    document.add_heading('Technical Implementation', 0)
    document.add_heading("Code", level=1)

    initial=os.getcwd()
    print(os.getcwd())
    y=input("y or n")
    if (y=="y"):
        for (root,dirs,files) in os.walk('.', topdown=True):
            print(root)
            for file in files:
                ys=input(f"include {file}?")
                if(ys=="y"):
                    add_heading(root+"/"+file)
                    read_code(root+"/"+file)
                    save_doc()
            dirscopy=dirs.copy()
            for d in dirs:
                inc=input(f"include dir {d}?")
                if(inc=="n"):
                    dirscopy.remove(d)
            [dirs.remove(d) for d in list(dirs) if not d in dirscopy]
            #dirs[:] = [d for d in dirs if (lambda : input(f"include dir {d}?")=="n") ]